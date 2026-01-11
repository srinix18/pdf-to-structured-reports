"""
Create missing DOCX files for sections that only have JSON files.
"""
import sys
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(Path(__file__).parent.parent))
from config.config import OUTPUT_DIR
import logging

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def create_docx_from_json(json_file, docx_file):
    """Create a DOCX file from existing JSON file"""
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    doc = Document()
    
    # Add title
    title = doc.add_heading(data['section'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    metadata = doc.add_paragraph()
    metadata.add_run(f"Company: {data['company']}\n").bold = True
    metadata.add_run(f"Year: {data['year']}\n").bold = True
    if 'detected_heading' in data:
        metadata.add_run(f"Detected Heading: {data['detected_heading']}\n").italic = True
    
    doc.add_paragraph()  # Empty line
    
    # Add section content
    paragraphs = data['text'].split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            p = doc.add_paragraph(para_text.strip())
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
    
    doc.save(str(docx_file))
    logger.info(f"  ✓ Created {docx_file.name}")


def main():
    logger.info("=" * 80)
    logger.info("Creating Missing DOCX Files from JSON")
    logger.info("=" * 80)
    
    missing = []
    
    # Find JSON files without corresponding DOCX
    for company_dir in OUTPUT_DIR.iterdir():
        if not company_dir.is_dir():
            continue
            
        for year_dir in company_dir.iterdir():
            if not year_dir.is_dir():
                continue
                
            sections_dir = year_dir / "sections"
            if not sections_dir.exists():
                continue
            
            # Check both section types
            for section_name in ["mdna", "letter_to_stakeholders"]:
                json_file = sections_dir / f"{section_name}.json"
                docx_file = sections_dir / f"{section_name}.docx"
                
                if json_file.exists() and not docx_file.exists():
                    missing.append({
                        "company": company_dir.name,
                        "year": year_dir.name,
                        "json_file": json_file,
                        "docx_file": docx_file,
                        "section": section_name
                    })
    
    if not missing:
        logger.info("No missing DOCX files found. All sections have both JSON and DOCX.")
        return
    
    logger.info(f"\nFound {len(missing)} sections with JSON but no DOCX:")
    mdna_count = sum(1 for m in missing if m['section'] == 'mdna')
    letter_count = sum(1 for m in missing if m['section'] == 'letter_to_stakeholders')
    logger.info(f"  MD&A: {mdna_count}")
    logger.info(f"  Letter: {letter_count}")
    
    response = input(f"\nCreate DOCX files for these {len(missing)} sections? (y/n): ")
    if response.lower() != 'y':
        logger.info("Aborted by user.")
        return
    
    logger.info("\nCreating DOCX files...")
    success = 0
    errors = 0
    
    for item in missing:
        logger.info(f"[{success + errors + 1}/{len(missing)}] {item['company']} - {item['year']} ({item['section']})")
        try:
            create_docx_from_json(item['json_file'], item['docx_file'])
            success += 1
        except Exception as e:
            logger.error(f"  ✗ Error: {e}")
            errors += 1
    
    logger.info("\n" + "=" * 80)
    logger.info(f"Summary:")
    logger.info(f"  ✓ Created: {success}")
    logger.info(f"  ✗ Errors: {errors}")
    logger.info("=" * 80)


if __name__ == "__main__":
    main()
