"""
Re-extract MD&A and/or Letter to Stakeholders from existing DOCX files.
Only extracts sections that are missing - preserves existing extractions.
"""
import sys
import json
from pathlib import Path
from docx import Document

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from config.config import OUTPUT_DIR
from pipeline.section_metadata import SectionType, SECTION_KEYWORDS
import logging
import re

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def find_reports_missing_sections():
    """Find all reports missing MD&A and/or Letter to Stakeholders"""
    missing = []
    
    for company_dir in OUTPUT_DIR.iterdir():
        if not company_dir.is_dir():
            continue
            
        for year_dir in company_dir.iterdir():
            if not year_dir.is_dir():
                continue
                
            sections_dir = year_dir / "sections"
            mdna_file = sections_dir / "mdna.json"
            letter_file = sections_dir / "letter_to_stakeholders.json"
            docx_file = year_dir / "report.docx"
            
            # Check if DOCX exists and has sections dir
            if not docx_file.exists() or not sections_dir.exists():
                continue
            
            # Check what's missing
            missing_sections = []
            if not mdna_file.exists():
                missing_sections.append("mdna")
            if not letter_file.exists():
                missing_sections.append("letter")
            
            if missing_sections:
                missing.append({
                    "company": company_dir.name,
                    "year": year_dir.name,
                    "docx_path": docx_file,
                    "sections_dir": sections_dir,
                    "missing_sections": missing_sections
                })
    
    return missing


def read_docx_with_structure(docx_path):
    """Read DOCX and extract text with heading information"""
    doc = Document(docx_path)
    
    # Get all keywords to look for
    all_keywords = set()
    for keywords in SECTION_KEYWORDS.values():
        all_keywords.update(keywords)
    
    sections = []
    current_section = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or text == "=== PAGE BREAK ===":
            continue
        
        # Skip page number headers
        if text.startswith("Page ") and len(text) < 10:
            continue
        
        # Check if this paragraph contains section keywords (potential heading)
        text_lower = text.lower()
        text_normalized = re.sub(r'[^\w\s]', ' ', text_lower)
        text_normalized = ' '.join(text_normalized.split())
        
        is_potential_heading = False
        matched_keyword = None
        for keyword in all_keywords:
            if keyword in text_normalized and len(text) < 300:
                is_potential_heading = True
                matched_keyword = keyword
                break
        
        if is_potential_heading:
            # Save previous section if it has content
            if current_section and len(current_section["text"]) > 100:  # At least 100 chars
                sections.append(current_section)
            
            # Start new section
            current_section = {"heading": text, "text": text + "\n", "keyword": matched_keyword}
        elif current_section:
            current_section["text"] += text + "\n"
    
    # Add last section
    if current_section and len(current_section["text"]) > 100:
        sections.append(current_section)
    
    return sections


def find_section_in_docx(sections, section_type):
    """Find a section by matching keywords"""
    keywords = SECTION_KEYWORDS.get(section_type, [])
    
    best_match = None
    best_score = 0
    
    for i, section in enumerate(sections):
        # Check if any of this section type's keywords match
        matched_keyword = section.get("keyword", "")
        
        if matched_keyword in keywords:
            # Calculate match score (longer matches = better)
            score = len(matched_keyword)
            if score > best_score:
                best_score = score
                best_match = {
                    "section_index": i,
                    "heading": section["heading"],
                    "text": section["text"],
                    "keyword_matched": matched_keyword
                }
    
    return best_match


def save_section_json(section_data, output_file, company, year, section_type):
    """Save extracted section to JSON file"""
    section_name_map = {
        SectionType.MDNA: "MD&A",
        SectionType.LETTER_TO_STAKEHOLDERS: "Letter To Stakeholders"
    }
    
    output_data = {
        "company": company,
        "year": year,
        "section": section_name_map.get(section_type, section_type.value),
        "start_page": 1,  # From DOCX we don't have exact page numbers
        "end_page": 1,
        "confidence": 1.0,
        "detected_heading": section_data["heading"],
        "keyword_matched": section_data["keyword_matched"],
        "text": section_data["text"],
        "character_count": len(section_data["text"]),
        "extraction_method": "docx_keyword_matching"
    }
    
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, indent=2, ensure_ascii=False)
    
    return output_file


def save_section_docx(section_data, output_file, company, year, section_type):
    """Save extracted section to DOCX file"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Add title
    section_name_map = {
        SectionType.MDNA: "Management Discussion & Analysis",
        SectionType.LETTER_TO_STAKEHOLDERS: "Letter To Stakeholders"
    }
    title = doc.add_heading(section_name_map.get(section_type, section_type.value), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    metadata = doc.add_paragraph()
    metadata.add_run(f"Company: {company}\n").bold = True
    metadata.add_run(f"Year: {year}\n").bold = True
    metadata.add_run(f"Detected Heading: {section_data['heading']}\n").italic = True
    
    doc.add_paragraph()  # Empty line
    
    # Add section content
    # Split by paragraphs and add each one
    paragraphs = section_data['text'].split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            p = doc.add_paragraph(para_text.strip())
            # Set font
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
    
    doc.save(str(output_file))
    return output_file


def reextract_sections_for_report(report_info):
    """
    Re-extract missing sections from DOCX file.
    
    Args:
        report_info: Dict with company, year, docx_path, sections_dir, missing_sections
        
    Returns:
        Dict with extraction results
    """
    company = report_info["company"]
    year = report_info["year"]
    docx_path = report_info["docx_path"]
    sections_dir = report_info["sections_dir"]
    missing = report_info["missing_sections"]
    
    logger.info(f"Extracting {', '.join(missing)} from DOCX")
    
    if not docx_path.exists():
        logger.warning(f"DOCX not found: {docx_path}")
        return {"success": False, "error": "docx_not_found"}
    
    try:
        # Read DOCX structure
        logger.info(f"  Reading DOCX file...")
        sections = read_docx_with_structure(docx_path)
        logger.info(f"  Found {len(sections)} sections in DOCX")
        
        # Extract only missing sections
        results = {}
        for section_type_str in missing:
            section_type = SectionType.MDNA if section_type_str == "mdna" else SectionType.LETTER_TO_STAKEHOLDERS
            
            # Find section in DOCX
            match = find_section_in_docx(sections, section_type)
            
            if match:
                # Save section in both JSON and DOCX formats
                filename_base = "letter_to_stakeholders" if section_type_str == "letter" else section_type_str
                json_file = sections_dir / f"{filename_base}.json"
                docx_file = sections_dir / f"{filename_base}.docx"
                
                save_section_json(match, json_file, company, year, section_type)
                save_section_docx(match, docx_file, company, year, section_type)
                
                results[section_type_str] = "extracted"
                logger.info(f"  ✓ Extracted {section_type_str} (matched: '{match['keyword_matched']}')")
            else:
                results[section_type_str] = "not_found"
                logger.info(f"  ✗ {section_type_str} not found")
        
        return {"success": True, "results": results}
            
    except Exception as e:
        logger.error(f"Error processing {company} - {year}: {e}", exc_info=True)
        return {"success": False, "error": str(e)}


def main():
    """Main function to re-extract missing sections"""
    logger.info("=" * 80)
    logger.info("Re-extracting Missing Sections from DOCX Files")
    logger.info("=" * 80)
    
    # Find reports with missing sections
    logger.info("Scanning for reports with missing sections...")
    missing_reports = find_reports_missing_sections()
    
    if not missing_reports:
        logger.info("All reports have both MD&A and Letter extracted. Nothing to do.")
        return
    
    # Count what's missing
    mdna_missing = sum(1 for r in missing_reports if "mdna" in r["missing_sections"])
    letter_missing = sum(1 for r in missing_reports if "letter" in r["missing_sections"])
    
    logger.info(f"\nFound:")
    logger.info(f"  Reports missing MD&A: {mdna_missing}")
    logger.info(f"  Reports missing Letter: {letter_missing}")
    logger.info(f"  Total reports to process: {len(missing_reports)}")
    
    # Ask for confirmation
    print(f"\nAbout to re-extract sections for {len(missing_reports)} reports from DOCX files.")
    print("This will only extract sections that are missing.")
    print("Existing extractions will not be affected.")
    response = input("\nProceed? (y/n): ")
    
    if response.lower() != 'y':
        logger.info("Aborted by user.")
        return
    
    # Re-extract
    import time
    start_time = time.time()
    
    logger.info("\nStarting re-extraction...")
    stats = {
        "mdna_extracted": 0,
        "mdna_not_found": 0,
        "letter_extracted": 0,
        "letter_not_found": 0,
        "errors": 0
    }
    
    for i, report_info in enumerate(missing_reports, 1):
        # Progress header with stats
        progress_pct = (i / len(missing_reports)) * 100
        elapsed = time.time() - start_time
        avg_time = elapsed / i if i > 0 else 0
        remaining = avg_time * (len(missing_reports) - i)
        
        logger.info(f"\n{'='*80}")
        logger.info(f"[{i}/{len(missing_reports)}] ({progress_pct:.1f}%) {report_info['company']} - {report_info['year']}")
        logger.info(f"  Missing: {', '.join(report_info['missing_sections'])}")
        logger.info(f"  Running stats - MD&A: ✓{stats['mdna_extracted']} ✗{stats['mdna_not_found']} | Letter: ✓{stats['letter_extracted']} ✗{stats['letter_not_found']}")
        if i > 1:
            logger.info(f"  Time: {elapsed/60:.1f}min elapsed, ~{remaining/60:.1f}min remaining")
        
        try:
            result = reextract_sections_for_report(report_info)
            
            if result["success"]:
                for section_type, status in result["results"].items():
                    stats[f"{section_type}_{status}"] += 1
            else:
                stats["errors"] += 1
                
        except KeyboardInterrupt:
            logger.warning("\n\n" + "="*80)
            logger.warning("INTERRUPTED BY USER - Saving progress...")
            logger.warning("="*80)
            break
        except Exception as e:
            logger.error(f"Unexpected error: {e}", exc_info=True)
            stats["errors"] += 1
    
    # Summary
    logger.info("\n" + "=" * 80)
    logger.info("Re-extraction Summary")
    logger.info("=" * 80)
    logger.info(f"Total reports processed: {i}")
    logger.info(f"\nMD&A:")
    logger.info(f"  ✓ Extracted: {stats['mdna_extracted']}")
    logger.info(f"  ✗ Not found: {stats['mdna_not_found']}")
    logger.info(f"\nLetter to Stakeholders:")
    logger.info(f"  ✓ Extracted: {stats['letter_extracted']}")
    logger.info(f"  ✗ Not found: {stats['letter_not_found']}")
    logger.info(f"\nErrors: {stats['errors']}")
    
    # Calculate new rates
    total_reports = 239
    current_mdna = 152
    current_letter = 61
    
    new_mdna = current_mdna + stats['mdna_extracted']
    new_letter = current_letter + stats['letter_extracted']
    
    logger.info(f"\nProjected Extraction Rates:")
    logger.info(f"  MD&A: {new_mdna}/{total_reports} ({new_mdna/total_reports*100:.1f}%) - was {current_mdna/total_reports*100:.1f}%")
    logger.info(f"  Letter: {new_letter}/{total_reports} ({new_letter/total_reports*100:.1f}%) - was {current_letter/total_reports*100:.1f}%")


if __name__ == "__main__":
    main()
