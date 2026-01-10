"""
Module for exporting processed data to various formats.
"""
import logging
import json
from pathlib import Path
from typing import List, Dict, Any
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pipeline.extract_text import PageText
from config.config import OUTPUT_DIR

logger = logging.getLogger(__name__)


def create_output_directory(company_name: str, year: str) -> Path:
    """
    Create output directory structure for a company/year.
    
    Args:
        company_name: Name of the company
        year: Report year
        
    Returns:
        Path to the output directory
    """
    output_path = OUTPUT_DIR / company_name / year
    output_path.mkdir(parents=True, exist_ok=True)
    
    # Create subdirectory for tables
    tables_path = output_path / "tables"
    tables_path.mkdir(exist_ok=True)
    
    logger.info(f"Created output directory: {output_path}")
    return output_path


def export_to_docx(
    pages: List[PageText],
    sections: List,
    output_path: Path,
    company_name: str,
    year: str
) -> Path:
    """
    Export extracted text to DOCX format - PAGE BY PAGE for best readability.
    Simpler structure that preserves original document flow.
    
    Args:
        pages: List of PageText objects
        sections: List of Section objects (not used in this version)
        output_path: Directory to save the file
        company_name: Company name
        year: Report year
        
    Returns:
        Path to the created DOCX file
    """
    logger.info(f"Exporting to DOCX: {company_name} {year}")
    
    doc = Document()
    
    # Set document title
    title = doc.add_heading(f'{company_name} - Annual Report {year}', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total Pages: {len(pages)}")
    doc.add_paragraph(f"Extraction Method: {'OCR' if pages[0].method == 'ocr' else 'Direct Text Extraction'}")
    doc.add_page_break()
    
    # Export page by page for maximum readability and preservation of document flow
    logger.info(f"Adding {len(pages)} pages to document")
    
    for page in pages:
        # Add page marker
        page_heading = doc.add_heading(f'Page {page.page_number}', level=2)
        page_heading.style.font.size = Pt(12)
        page_heading.style.font.bold = True
        
        if page.text.strip():
            # Split into paragraphs - preserve original line breaks for layout
            paragraphs = page.text.split('\n\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Keep single line breaks within paragraphs for formatting preservation
                    para = doc.add_paragraph(para_text.strip())
                    para.style.font.size = Pt(10)
                    para.style.font.name = 'Calibri'
        else:
            doc.add_paragraph("[No text on this page]")
        
        # Add small spacing between pages
        doc.add_paragraph()
    
    # Save document
    docx_path = output_path / "report.docx"
    doc.save(str(docx_path))
    
    logger.info(f"DOCX saved to: {docx_path}")
    return docx_path


def export_tables_to_csv(
    tables: List,
    output_path: Path
) -> List[Path]:
    """
    Export extracted tables to CSV files.
    
    Args:
        tables: List of ExtractedTable objects
        output_path: Directory to save the files
        
    Returns:
        List of paths to created CSV files
    """
    logger.info(f"Exporting {len(tables)} tables to CSV")
    
    tables_dir = output_path / "tables"
    saved_files = []
    
    for idx, table in enumerate(tables):
        # Generate filename
        table_name = f"table_{idx+1}_page_{table.page_number}.csv"
        csv_path = tables_dir / table_name
        
        try:
            # Clean the dataframe
            df = table.dataframe.copy()
            
            # Save to CSV
            df.to_csv(csv_path, index=False)
            saved_files.append(csv_path)
            
            logger.debug(f"Saved table to: {csv_path}")
            
        except Exception as e:
            logger.error(f"Error saving table {idx+1}: {e}")
    
    logger.info(f"Saved {len(saved_files)} tables to CSV")
    return saved_files


def export_metadata_to_json(
    pdf_info: Dict[str, Any],
    pdf_type: str,
    pages: List[PageText],
    sections: List,
    tables: List,
    output_path: Path,
    company_name: str,
    year: str,
    financial_data: Dict[str, Any] = None
) -> Path:
    """
    Export metadata and processing information to JSON.
    
    Args:
        pdf_info: PDF file information
        pdf_type: Type of PDF (text/scanned)
        pages: List of PageText objects
        sections: List of Section objects
        tables: List of ExtractedTable objects
        output_path: Directory to save the file
        company_name: Company name
        year: Report year
        financial_data: Extracted financial data (optional)
        
    Returns:
        Path to the created JSON file
    """
    logger.info("Exporting metadata to JSON")
    
    metadata = {
        "company": company_name,
        "year": year,
        "processing_date": datetime.now().isoformat(),
        "pdf_info": pdf_info,
        "pdf_type": pdf_type,
        "statistics": {
            "total_pages": len(pages),
            "total_characters": sum(page.char_count for page in pages),
            "total_sections": len(sections),
            "total_tables": len(tables)
        },
        "sections": [
            {
                "title": section.title,
                "level": section.level,
                "start_page": section.start_page,
                "end_page": section.end_page,
                "character_count": len(section.content)
            }
            for section in sections
        ],
        "tables": [
            {
                "index": idx,
                "page": table.page_number,
                "rows": table.rows,
                "columns": table.cols,
                "extraction_method": table.method
            }
            for idx, table in enumerate(tables)
        ],
        "pages": [
            {
                "page_number": page.page_number,
                "character_count": page.char_count,
                "extraction_method": page.method
            }
            for page in pages
        ]
    }
    
    # Add financial data if available
    if financial_data:
        metadata["financial_data"] = financial_data
    
    json_path = output_path / "metadata.json"
    
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(metadata, f, indent=2, ensure_ascii=False)
    
    logger.info(f"Metadata saved to: {json_path}")
    return json_path


def export_all(
    pdf_info: Dict[str, Any],
    pdf_type: str,
    pages: List[PageText],
    sections: List,
    tables: List,
    company_name: str,
    year: str,
    financial_data: Dict[str, Any] = None
) -> Dict[str, Any]:
    """
    Export all data in all formats.
    
    Args:
        pdf_info: PDF file information
        pdf_type: Type of PDF (text/scanned)
        pages: List of PageText objects
        sections: List of Section objects
        tables: List of ExtractedTable objects
        company_name: Company name
        year: Report year
        financial_data: Extracted financial data (optional)
        
    Returns:
        Dictionary with paths to all created files
    """
    logger.info(f"Starting export for {company_name} {year}")
    
    # Create output directory
    output_path = create_output_directory(company_name, year)
    
    results = {
        "output_directory": str(output_path),
        "files_created": []
    }
    
    # Export to DOCX
    try:
        docx_path = export_to_docx(pages, sections, output_path, company_name, year)
        results["docx"] = str(docx_path)
        results["files_created"].append(str(docx_path))
    except Exception as e:
        logger.error(f"Error exporting to DOCX: {e}")
        results["docx_error"] = str(e)
    
    # Export tables to CSV
    try:
        csv_paths = export_tables_to_csv(tables, output_path)
        results["csv_files"] = [str(p) for p in csv_paths]
        results["files_created"].extend([str(p) for p in csv_paths])
    except Exception as e:
        logger.error(f"Error exporting tables: {e}")
        results["csv_error"] = str(e)
    
    # Export metadata to JSON
    try:
        json_path = export_metadata_to_json(
            pdf_info, pdf_type, pages, sections, tables,
            output_path, company_name, year, financial_data
        )
        results["metadata"] = str(json_path)
        results["files_created"].append(str(json_path))
    except Exception as e:
        logger.error(f"Error exporting metadata: {e}")
        results["metadata_error"] = str(e)
    
    logger.info(f"Export completed. Created {len(results['files_created'])} files")
    return results


def create_summary_report(company_results: List[Dict[str, Any]]) -> Path:
    """
    Create a summary report for all processed companies.
    
    Args:
        company_results: List of processing results for each company
        
    Returns:
        Path to the summary report
    """
    summary_path = OUTPUT_DIR / "processing_summary.json"
    
    summary = {
        "processing_date": datetime.now().isoformat(),
        "total_companies": len(company_results),
        "companies": company_results
    }
    
    with open(summary_path, 'w', encoding='utf-8') as f:
        json.dump(summary, f, indent=2, ensure_ascii=False)
    
    logger.info(f"Summary report saved to: {summary_path}")
    return summary_path
