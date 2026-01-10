"""
Section content extractor - extracts section content from already-processed text.
"""
import logging
import json
from pathlib import Path
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

from pipeline.extract_text import PageText
from pipeline.section_metadata import SectionBoundary, SectionContent, SectionType
from pipeline.section_boundary_detector import SectionBoundaryDetector

logger = logging.getLogger(__name__)


class SectionContentExtractor:
    """
    Extracts section content from already-processed page text using detected boundaries.
    """
    
    def __init__(self, pages: List[PageText], output_dir: Path):
        """
        Initialize extractor with processed page text.
        
        Args:
            pages: List of PageText objects from main pipeline
            output_dir: Directory to save section files
        """
        self.pages = pages
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
    def extract_section(
        self, 
        boundary: SectionBoundary
    ) -> Optional[SectionContent]:
        """
        Extract section content using detected boundary.
        
        Args:
            boundary: SectionBoundary with page range
            
        Returns:
            SectionContent with extracted text
        """
        if boundary.end_page is None:
            logger.warning(
                f"Cannot extract {boundary.section_type.value}: end page not determined"
            )
            return None
        
        logger.info(
            f"Extracting {boundary.section_type.value} content from pages "
            f"{boundary.start_page}-{boundary.end_page}"
        )
        
        # Extract pages in range
        section_pages = [
            p for p in self.pages 
            if boundary.start_page <= p.page_number <= boundary.end_page
        ]
        
        if not section_pages:
            logger.warning(
                f"No pages found in range {boundary.start_page}-{boundary.end_page}"
            )
            return None
        
        # Combine text from all pages
        section_text = "\n\n".join([p.text for p in section_pages])
        
        content = SectionContent(
            section_type=boundary.section_type,
            start_page=boundary.start_page,
            end_page=boundary.end_page,
            text=section_text,
            character_count=len(section_text),
            page_count=len(section_pages)
        )
        
        logger.info(
            f"Extracted {content.character_count:,} characters across "
            f"{content.page_count} pages"
        )
        
        return content
    
    def export_section_to_docx(
        self,
        content: SectionContent,
        company_name: str,
        year: str
    ) -> Path:
        """
        Export section content to separate DOCX file.
        
        Args:
            content: SectionContent to export
            company_name: Company name for title
            year: Report year
            
        Returns:
            Path to created DOCX file
        """
        section_name = content.section_type.value.replace('_', ' ').title()
        filename = f"{content.section_type.value}.docx"
        output_path = self.output_dir / filename
        
        logger.info(f"Exporting {section_name} to {output_path}")
        
        doc = Document()
        
        # Add title
        title = doc.add_heading(
            f'{company_name} - {section_name} ({year})', 
            level=0
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Pages: {content.start_page}-{content.end_page}")
        doc.add_paragraph(f"Character Count: {content.character_count:,}")
        doc.add_page_break()
        
        # Add content
        # Split into paragraphs for better formatting
        paragraphs = content.text.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.style.font.size = Pt(11)
        
        # Save document
        doc.save(output_path)
        logger.info(f"Saved {output_path}")
        
        return output_path
    
    def export_section_metadata(
        self,
        boundaries: Dict[str, Optional[SectionBoundary]],
        contents: Dict[str, Optional[SectionContent]]
    ) -> Path:
        """
        Export section metadata to JSON file.
        
        Args:
            boundaries: Dictionary of detected boundaries
            contents: Dictionary of extracted contents
            
        Returns:
            Path to created JSON file
        """
        output_path = self.output_dir / "sections_metadata.json"
        
        logger.info(f"Exporting section metadata to {output_path}")
        
        metadata = {}
        
        for section_key in boundaries.keys():
            boundary = boundaries.get(section_key)
            content = contents.get(section_key)
            
            if boundary:
                metadata[section_key] = {
                    "boundary": boundary.to_dict(),
                    "content_stats": content.to_dict() if content else None,
                    "extracted": content is not None
                }
            else:
                metadata[section_key] = {
                    "boundary": None,
                    "content_stats": None,
                    "extracted": False,
                    "note": "Section not found in document"
                }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Saved metadata to {output_path}")
        
        return output_path


def extract_sections_from_pdf(
    pdf_path: Path,
    pages: List[PageText],
    output_dir: Path,
    company_name: str,
    year: str
) -> Dict[str, Path]:
    """
    Complete section extraction workflow.
    
    Args:
        pdf_path: Path to source PDF
        pages: Processed PageText objects from main pipeline
        output_dir: Directory for section outputs
        company_name: Company name
        year: Report year
        
    Returns:
        Dictionary mapping section type to output file paths
    """
    logger.info("="*80)
    logger.info("Starting section extraction")
    logger.info("="*80)
    
    output_files = {}
    
    try:
        # Step 1: Detect boundaries from PDF layout
        logger.info("Step 1: Detecting section boundaries from PDF layout...")
        detector = SectionBoundaryDetector(pdf_path)
        boundaries = detector.detect_section_boundaries()
        
        # Step 2: Extract content from processed text
        logger.info("Step 2: Extracting section content from processed text...")
        extractor = SectionContentExtractor(pages, output_dir)
        
        contents = {}
        for section_key, boundary in boundaries.items():
            if boundary:
                content = extractor.extract_section(boundary)
                contents[section_key] = content
                
                # Export to DOCX
                if content:
                    docx_path = extractor.export_section_to_docx(
                        content, company_name, year
                    )
                    output_files[section_key] = docx_path
            else:
                contents[section_key] = None
                logger.info(f"Section {section_key} not found")
        
        # Step 3: Export metadata
        logger.info("Step 3: Exporting section metadata...")
        metadata_path = extractor.export_section_metadata(boundaries, contents)
        output_files['metadata'] = metadata_path
        
        logger.info("="*80)
        logger.info(f"Section extraction complete. Created {len(output_files)} files")
        logger.info("="*80)
        
    except Exception as e:
        logger.error(f"Error during section extraction: {e}", exc_info=True)
    
    return output_files
